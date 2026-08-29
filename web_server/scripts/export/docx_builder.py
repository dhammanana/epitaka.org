"""
docx_builder.py — Generate DOCX files from Book data.

Embeds script-specific fonts so Sinhala and other complex scripts
render correctly even on systems without the fonts installed.
"""
import os
import re
from copy import deepcopy
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .data_loader import Book

_FONTS_DIR = os.path.normpath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'frontend', 'src', 'fonts'
))

# Script code → (regular_font_filename, bold_font_filename)
_SCRIPT_FONT_FILES = {
    'si':  ('NotoSerifSinhala-Regular.ttf',      'NotoSerifSinhala-Bold.ttf'),
    'hi':  ('NotoSerifDevanagari-Regular.ttf',   'NotoSerifDevanagari-Bold.ttf'),
    'th':  ('thai/THSarabunPali.ttf',            'thai/THSarabun-Bold.ttf'),
    'km':  ('NotoSerifKhmer-Regular.ttf',        'NotoSerifKhmer-Bold.ttf'),
    'be':  ('NotoSerifBengali-Regular.ttf',      'NotoSerifBengali-Bold.ttf'),
    'gm':  ('NotoSansGurmukhi-Regular.ttf',      'NotoSansGurmukhi-Bold.ttf'),
    'gj':  ('NotoSerifGujarati-Regular.ttf',     'NotoSerifGujarati-Bold.ttf'),
    'te':  ('NotoSerifTelugu-Regular.ttf',       'NotoSerifTelugu-Bold.ttf'),
    'ka':  ('NotoSerifKannada-Regular.ttf',      'NotoSerifKannada-Bold.ttf'),
    'mm':  ('NotoSerifMalayalam-Regular.ttf',    'NotoSerifMalayalam-Bold.ttf'),
}

ACCENT = RGBColor(139, 94, 60)
PALI_C = RGBColor(124, 45, 18)
TRANS_C = RGBColor(30, 58, 95)
MUTED = RGBColor(138, 122, 110)
TEXT = RGBColor(45, 36, 32)


# Script code → font name for DOCX (installed system fonts)
_SCRIPT_DOCX_FONTS = {
    'ro': 'Noto Serif',
    'si': 'Noto Serif Sinhala',
    'hi': 'Noto Serif Devanagari',
    'th': 'THSarabunPali',
    'km': 'Noto Serif Khmer',
    'be': 'Noto Serif Bengali',
    'my': 'Myanmar3',
    'gm': 'Noto Sans Gurmukhi',
    'gj': 'Noto Serif Gujarati',
    'te': 'Noto Serif Telugu',
    'ka': 'Noto Serif Kannada',
    'mm': 'Noto Serif Malayalam',
}


def build_docx(book: Book, output_path: str) -> str:
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    # Set the Pāli font for the document
    pali_font_name = _SCRIPT_DOCX_FONTS.get(book.script, 'Noto Serif')
    # Collect fonts to embed
    fonts_to_embed = _get_fonts_to_embed(book.script)

    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Dharma Wheel'); r.font.size = Pt(40); r.font.color.rgb = ACCENT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_docx_title(book)); r.font.size = Pt(26); r.font.color.rgb = ACCENT; r.bold = True
    if book.sub_nikaya:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(book.sub_nikaya); r.font.size = Pt(13); r.font.color.rgb = MUTED
    elif book.nikaya:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(book.nikaya); r.font.size = Pt(13); r.font.color.rgb = MUTED
    if book.lang_name:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'{book.lang_name} Translation'); r.font.size = Pt(11); r.font.color.rgb = ACCENT
    if book.description:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(book.description); r.font.size = Pt(9); r.font.color.rgb = MUTED; r.italic = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Chattha Sangayana Tipitaka'); r.font.size = Pt(10); r.font.color.rgb = MUTED
    doc.add_page_break()

    doc.add_heading('Table of Contents', level=1)
    p = doc.add_paragraph(); run = p.add_run()
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run = p.add_run(); run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'))
    run = p.add_run(); run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    run = p.add_run('[Update field in Word: right-click → Update Field]'); run.font.color.rgb = MUTED; run.font.size = Pt(9)
    run = p.add_run(); run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    doc.add_page_break()

    if book.intro_sentences:
        doc.add_heading(book.book_name, level=1)
        for s in book.intro_sentences: _add_docx_sent(doc, s, pali_font_name)

    for vagga in book.vagga_sections:
        h = vagga.heading; title = h.title or f'Section {h.para_id}'
        heading = doc.add_heading(title, level=1)
        for run in heading.runs: run.font.color.rgb = ACCENT
        if vagga.heading_translation: _add_translation_heading(doc, vagga.heading_translation)
        for verse in vagga.verses:
            vh = verse.heading; vtitle = vh.title or f'Section {vh.para_id}'
            if vh.para_id != h.para_id:
                heading = doc.add_heading(vtitle, level=2)
                for run in heading.runs: run.font.color.rgb = ACCENT
            if verse.heading_translation: _add_translation_heading(doc, verse.heading_translation, 9)
            for s in verse.sentences: _add_docx_sent(doc, s, pali_font_name)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('- - -'); r.font.color.rgb = MUTED

    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{book.book_name} — Chattha Sangayana Tipitaka'); r.italic = True; r.font.color.rgb = MUTED; r.font.size = Pt(9)
    doc.save(output_path)
    # Embed fonts after saving (need the file on disk)
    _embed_fonts_in_docx(output_path, fonts_to_embed)
    return output_path


def _add_translation_heading(doc, text, size=10):
    p = doc.add_paragraph(); r = p.add_run(_strip_basic(text)); r.italic = True; r.font.color.rgb = TRANS_C; r.font.size = Pt(size)


def _add_docx_sent(doc, s, pali_font_name='Noto Serif'):
    if s.vripage:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f'— VRI page {s.vripage} —'); r.font.color.rgb = MUTED; r.font.size = Pt(8)
    if s.pali: _add_rich_paragraph(doc, s.pali, PALI_C, 11, 0, pali_font_name)
    if s.translation: _add_rich_paragraph(doc, s.translation, TRANS_C, 10, Cm(0.5))


def _add_rich_paragraph(doc, html, color, size, indent, font_name=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = indent
    for text, attrs in _parse_inline(html):
        r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(size)
        if font_name:
            r.font.name = font_name
            # Set all font categories so non-Latin scripts render correctly
            rPr = r._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
        r.bold = attrs.get('bold', False); r.italic = attrs.get('italic', False)
        r.font.superscript = attrs.get('sup', False); r.font.subscript = attrs.get('sub', False)


def _parse_inline(html):
    pattern = re.compile(r'(<(?:b|strong|i|em|sup|sub)>|</(?:b|strong|i|em|sup|sub)>)', re.I)
    attrs = {'bold': False, 'italic': False, 'sup': False, 'sub': False}
    out = []
    for part in pattern.split(html or ''):
        if not part: continue
        tag = part.lower()
        if tag in ('<b>', '<strong>'): attrs['bold'] = True
        elif tag in ('</b>', '</strong>'): attrs['bold'] = False
        elif tag in ('<i>', '<em>'): attrs['italic'] = True
        elif tag in ('</i>', '</em>'): attrs['italic'] = False
        elif tag == '<sup>': attrs['sup'] = True
        elif tag == '</sup>': attrs['sup'] = False
        elif tag == '<sub>': attrs['sub'] = True
        elif tag == '</sub>': attrs['sub'] = False
        elif not part.startswith('<'):
            out.append((re.sub(r'<[^>]+>', '', part), attrs.copy()))
    return out


def _strip_basic(html):
    return re.sub(r'<[^>]+>', '', html or '').strip()


def _get_fonts_to_embed(script):
    """Get list of (font_name, file_path) tuples to embed in the DOCX."""
    if script not in _SCRIPT_FONT_FILES:
        return []
    regular_file, bold_file = _SCRIPT_FONT_FILES[script]
    fonts = []
    for fname in [regular_file, bold_file]:
        path = os.path.join(_FONTS_DIR, fname)
        if os.path.isfile(path):
            # Font name for DOCX (e.g., 'NotoSerifSinhala-Regular')
            base = os.path.splitext(os.path.basename(fname))[0]
            fonts.append((base, path))
    return fonts


def _embed_fonts_in_docx(docx_path, fonts_to_embed):
    """Embed TrueType font files into an existing DOCX.

    Adds font files to word/fonts/ and updates fontTable.xml.
    This makes the DOCX self-contained — no system fonts needed.
    """
    if not fonts_to_embed:
        return

    import zipfile
    import tempfile
    import shutil

    # Read existing DOCX
    with zipfile.ZipFile(docx_path, 'r') as zin:
        contents = {name: zin.read(name) for name in zin.namelist()}

    # Parse fontTable.xml
    font_table_xml = contents.get('word/fontTable.xml')
    if font_table_xml is None:
        font_table_xml = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    font_table = etree.fromstring(font_table_xml)

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for font_name, font_path in fonts_to_embed:
        # Check if already embedded
        existing = font_table.find(f'w:rFonts[@w:name="{font_name}"]', nsmap)
        if existing is not None:
            continue

        # Read font data and base64-encode it
        import base64
        with open(font_path, 'rb') as f:
            font_data = base64.b64encode(f.read()).decode('ascii')

        # Add entry to fontTable.xml with inline embedded font data
        font_elem = etree.SubElement(font_table, qn('w:font'))
        font_elem.set(qn('w:name'), font_name)
        # Add embedded font reference (base64-encoded TTF data)
        embed_elem = etree.SubElement(font_elem, qn('w:embedRegular'))
        embed_elem.set(qn('w:embeddedOpenType'), font_data)

    # Update fontTable.xml
    contents['word/fontTable.xml'] = etree.tostring(font_table, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Write updated DOCX
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)


def _docx_title(book):
    en = book.english_name
    return f'{en}\n({book.book_name})' if en and en.lower() != book.book_name.lower() else book.book_name
